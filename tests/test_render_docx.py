"""Tests for src/render_docx.py.

Skips entirely if python-docx isn't installed locally (CI installs it
from requirements.txt). Tests open the rendered .docx with python-docx
and assert on structure rather than visual layout.

Run with: python -m unittest tests.test_render_docx
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

import tempfile
import unittest
from pathlib import Path

try:
    import docx as _docx  # noqa: F401
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


PROFILE = {
    "person": {
        "name": "Sandeep Sahasranamam",
        "email": "sandeep@example.com",
        "location": "Boston, MA",
        "linkedin": "linkedin.com/in/sandeep",
    },
    "experience": [
        {
            "id": "amazon",
            "company": "Amazon",
            "location": "Seattle, WA",
            "title": "Senior TPM",
            "start": "2018-01",
            "end": "2023-12",
            "achievements": [
                {"id": "gdpr", "metrics": ["50 engineers"]},
                {"id": "echo", "metrics": []},
            ],
        },
        {
            "id": "oldco",
            "company": "OldCo",
            "location": "Remote",
            "title": "PM",
            "start": "2015-01",
            "end": "2017-12",
            "achievements": [{"id": "unused_ach"}],
        },
    ],
    "skills": {
        "program_management": ["Roadmap definition", "Risk mitigation"],
        "technical": ["AWS", "Python", "GCP"],
    },
    "education": [
        {"school": "MIT", "degree": "MS Computer Science", "year": 2010},
    ],
}


def _bullet(ach_id: str, text: str, verified: bool = True) -> dict:
    return {"achievement_id": ach_id, "text": text, "verified": verified, "reason": ""}


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx not installed locally; CI runs these")
class TestRenderResume(unittest.TestCase):

    def _render(self, tailor_result: dict) -> Path:
        from src.render_docx import render_resume
        tmp = Path(tempfile.mkdtemp()) / "test.docx"
        return render_resume(PROFILE, tailor_result, tmp)

    def _all_text(self, path: Path) -> str:
        from docx import Document
        d = Document(str(path))
        return "\n".join(p.text for p in d.paragraphs)

    def test_file_is_created_and_valid(self):
        from docx import Document
        result = {
            "headline": "TPM leader",
            "bullets": [_bullet("gdpr", "Led GDPR program")],
            "gaps": [],
        }
        path = self._render(result)
        self.assertTrue(path.exists())
        # Opening it shouldn't raise — proves it's a valid .docx
        Document(str(path))

    def test_name_email_location_present(self):
        result = {"headline": "", "bullets": [_bullet("gdpr", "b")], "gaps": []}
        text = self._all_text(self._render(result))
        self.assertIn("Sandeep Sahasranamam", text)
        self.assertIn("sandeep@example.com", text)
        self.assertIn("Boston, MA", text)
        self.assertIn("linkedin.com/in/sandeep", text)

    def test_headline_when_provided(self):
        result = {"headline": "Distinguished TPM leader xyz", "bullets": [_bullet("gdpr", "b")], "gaps": []}
        text = self._all_text(self._render(result))
        self.assertIn("Distinguished TPM leader xyz", text)

    def test_headline_omitted_when_empty(self):
        result = {"headline": "", "bullets": [_bullet("gdpr", "GOOD BULLET")], "gaps": []}
        text = self._all_text(self._render(result))
        # GOOD BULLET appears, no headline placeholder text
        self.assertIn("GOOD BULLET", text)

    def test_unverified_bullets_excluded(self):
        result = {
            "headline": "",
            "bullets": [
                _bullet("gdpr", "GOOD BULLET", verified=True),
                _bullet("gdpr", "BAD HALLUCINATION", verified=False),
            ],
            "gaps": [],
        }
        text = self._all_text(self._render(result))
        self.assertIn("GOOD BULLET", text)
        self.assertNotIn("BAD HALLUCINATION", text)

    def test_experience_without_bullets_is_omitted(self):
        """OldCo has no tailored bullets — its block should not appear."""
        result = {"headline": "", "bullets": [_bullet("gdpr", "Led GDPR")], "gaps": []}
        text = self._all_text(self._render(result))
        self.assertIn("Amazon", text)
        self.assertNotIn("OldCo", text)

    def test_experience_metadata_rendered(self):
        result = {"headline": "", "bullets": [_bullet("gdpr", "Led GDPR")], "gaps": []}
        text = self._all_text(self._render(result))
        self.assertIn("Senior TPM", text)
        self.assertIn("Amazon", text)
        self.assertIn("Seattle, WA", text)
        self.assertIn("2018-01", text)
        self.assertIn("2023-12", text)

    def test_end_date_defaults_to_present(self):
        profile = {
            **PROFILE,
            "experience": [{
                "id": "now",
                "company": "Firely",
                "location": "Boston",
                "title": "Founder",
                "start": "2025-01",
                "end": None,
                "achievements": [{"id": "x"}],
            }],
        }
        from src.render_docx import render_resume
        from docx import Document
        tmp = Path(tempfile.mkdtemp()) / "test.docx"
        render_resume(profile, {"headline": "", "bullets": [_bullet("x", "Built thing")], "gaps": []}, tmp)
        text = "\n".join(p.text for p in Document(str(tmp)).paragraphs)
        self.assertIn("Present", text)

    def test_skills_section(self):
        result = {"headline": "", "bullets": [_bullet("gdpr", "b")], "gaps": []}
        text = self._all_text(self._render(result))
        self.assertIn("SKILLS", text)
        self.assertIn("Roadmap definition", text)
        self.assertIn("AWS", text)
        self.assertIn("Python", text)

    def test_education_section(self):
        result = {"headline": "", "bullets": [_bullet("gdpr", "b")], "gaps": []}
        text = self._all_text(self._render(result))
        self.assertIn("EDUCATION", text)
        self.assertIn("MIT", text)
        self.assertIn("MS Computer Science", text)
        self.assertIn("2010", text)

    def test_zero_bullets_still_produces_file(self):
        """If all bullets got dropped during verification, render skills+edu anyway."""
        result = {"headline": "Just a headline", "bullets": [], "gaps": ["X"]}
        path = self._render(result)
        self.assertTrue(path.exists())
        text = self._all_text(path)
        # No EXPERIENCE heading since no bullets
        self.assertNotIn("EXPERIENCE", text)
        # But skills + education still there
        self.assertIn("SKILLS", text)
        self.assertIn("EDUCATION", text)

    def test_ats_safe_no_tables(self):
        from docx import Document
        result = {"headline": "x", "bullets": [_bullet("gdpr", "b")], "gaps": []}
        path = self._render(result)
        d = Document(str(path))
        self.assertEqual(len(d.tables), 0, "ATS-safe templates use no tables")

    def test_ats_safe_no_images(self):
        from docx import Document
        result = {"headline": "x", "bullets": [_bullet("gdpr", "b")], "gaps": []}
        path = self._render(result)
        d = Document(str(path))
        self.assertEqual(len(d.inline_shapes), 0, "ATS-safe templates use no images")

    def test_ats_safe_single_section_no_columns(self):
        from docx import Document
        result = {"headline": "x", "bullets": [_bullet("gdpr", "b")], "gaps": []}
        path = self._render(result)
        d = Document(str(path))
        self.assertEqual(len(d.sections), 1)

    def test_bullet_with_unknown_achievement_id_dropped(self):
        """Defensive: if verifier let through a bullet whose achievement_id doesn't
        match any experience entry, render_docx should drop it rather than crash
        or render it under a phantom company."""
        result = {
            "headline": "",
            "bullets": [
                _bullet("gdpr", "Real bullet"),
                _bullet("nonexistent_id", "Phantom bullet"),
            ],
            "gaps": [],
        }
        text = self._all_text(self._render(result))
        self.assertIn("Real bullet", text)
        self.assertNotIn("Phantom bullet", text)

    def test_uses_calibri_font(self):
        from docx import Document
        result = {"headline": "", "bullets": [_bullet("gdpr", "b")], "gaps": []}
        path = self._render(result)
        d = Document(str(path))
        # The Normal style should be Calibri
        self.assertEqual(d.styles["Normal"].font.name, "Calibri")


if __name__ == "__main__":
    unittest.main()
