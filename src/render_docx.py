"""ATS-safe .docx rendering for tailored resumes.

Uses python-docx. Hard constraints (per PRD "ATS compatibility" section):
  - Single column
  - No tables anywhere (some templates use tables for layout — broken in
    Workday, Greenhouse, Lever, Taleo, iCIMS)
  - No images, no inline shapes
  - No headers or footers
  - Calibri throughout (name 16pt, section heading 13pt, body 11pt)
  - Bullets use Word's built-in "List Bullet" style

Tests (test_render_docx.py) read the rendered file back via python-docx
and assert the structure. ATS parse quality is a manual jobscan.co check.

The renderer takes:
  - profile (dict)         — full profile.json
  - tailor_result (dict)   — output of src/tailor.tailor() as asdict()
  - out_path (Path)        — where to write the .docx

Only bullets with verified=True are rendered. Unverified bullets are
silently skipped (the tailor module already logged/dropped them).
"""
from __future__ import annotations
import logging
from pathlib import Path

log = logging.getLogger(__name__)

FONT_NAME = "Calibri"
NAME_PT = 16
HEADING_PT = 13
BODY_PT = 11
CONTACT_PT = 10


def render_resume(profile: dict, tailor_result: dict, out_path: str | Path) -> Path:
    """Render the tailored resume to out_path. Returns the written path."""
    # Lazy import so test discovery and CI without python-docx don't blow up.
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    _set_page_margins(doc, Inches)
    _set_default_font(doc, Pt)

    person = profile.get("person", {}) or {}
    _add_name(doc, person.get("name", ""), Pt)
    _add_contact_line(doc, person, Pt)
    _add_headline(doc, tailor_result.get("headline", ""), Pt)

    verified_bullets = [
        b for b in tailor_result.get("bullets", []) or []
        if b.get("verified") and b.get("text")
    ]
    if verified_bullets:
        bullets_by_exp = _group_bullets_by_experience(profile, verified_bullets)
        if bullets_by_exp:
            _add_section_heading(doc, "EXPERIENCE", Pt)
            for exp in profile.get("experience", []) or []:
                exp_id = exp.get("id")
                if exp_id not in bullets_by_exp:
                    continue
                _add_experience_block(doc, exp, bullets_by_exp[exp_id], Pt)

    skills = profile.get("skills", {}) or {}
    if any(skills.values()):
        _add_section_heading(doc, "SKILLS", Pt)
        for category, items in skills.items():
            if not items:
                continue
            label = category.replace("_", " ").title()
            line = f"{label}: {', '.join(items)}"
            _add_body_line(doc, line, Pt)

    education = profile.get("education", []) or []
    rendered_edu = [
        " | ".join(filter(None, [
            e.get("school"),
            e.get("degree"),
            str(e.get("year")) if e.get("year") else "",
        ]))
        for e in education if isinstance(e, dict) and e.get("school")
    ]
    if rendered_edu:
        _add_section_heading(doc, "EDUCATION", Pt)
        for line in rendered_edu:
            _add_body_line(doc, line, Pt)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _set_page_margins(doc, Inches) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def _set_default_font(doc, Pt) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_PT)


def _add_name(doc, name: str, Pt) -> None:
    if not name:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(NAME_PT)


def _add_contact_line(doc, person: dict, Pt) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    parts = [person.get("email"), person.get("location"), person.get("linkedin")]
    line = " · ".join(p for p in parts if p)
    if not line:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = FONT_NAME
    run.font.size = Pt(CONTACT_PT)


def _add_headline(doc, headline: str, Pt) -> None:
    if not headline:
        return
    p = doc.add_paragraph()
    run = p.add_run(headline)
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_PT)


def _add_section_heading(doc, text: str, Pt) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(HEADING_PT)


def _add_experience_block(doc, exp: dict, bullets: list[dict], Pt) -> None:
    # Line 1 (bold): Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(exp.get("title", "") or "")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = Pt(BODY_PT)

    # Line 2 (regular, smaller): Company · Location · Start – End
    sub_parts = [exp.get("company"), exp.get("location")]
    date_part = _format_date_range(exp.get("start"), exp.get("end"))
    if date_part:
        sub_parts.append(date_part)
    sub_line = " · ".join(p for p in sub_parts if p)
    if sub_line:
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(sub_line)
        sub_run.font.name = FONT_NAME
        sub_run.font.size = Pt(CONTACT_PT)

    # Bullets
    for b in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        # add_paragraph(style=...) may create the paragraph with default run; clear and add ours
        for r in list(bp.runs):
            r.text = ""
        run = bp.add_run(b["text"])
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_PT)


def _add_body_line(doc, text: str, Pt) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_PT)


def _format_date_range(start: str | None, end: str | None) -> str:
    if not start and not end:
        return ""
    start_s = start or ""
    end_s = end or "Present"
    return f"{start_s} – {end_s}"


def _group_bullets_by_experience(profile: dict, verified_bullets: list[dict]) -> dict[str, list[dict]]:
    """Map experience.id -> list of bullets that reference an achievement in that experience.

    Bullets whose achievement_id doesn't match any profile experience are dropped
    here (defensive — verifier should have caught this, but we don't trust input).
    """
    ach_to_exp: dict[str, str] = {}
    for exp in profile.get("experience", []) or []:
        exp_id = exp.get("id")
        if not exp_id:
            continue
        for ach in exp.get("achievements", []) or []:
            ach_id = ach.get("id")
            if ach_id:
                ach_to_exp[ach_id] = exp_id

    grouped: dict[str, list[dict]] = {}
    for b in verified_bullets:
        ach_id = b.get("achievement_id")
        exp_id = ach_to_exp.get(ach_id)
        if not exp_id:
            log.warning("bullet references unknown achievement_id %r — dropping", ach_id)
            continue
        grouped.setdefault(exp_id, []).append(b)
    return grouped
